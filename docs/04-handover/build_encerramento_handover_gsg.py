from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
OUT_FILE = OUT_DIR / "encerramento_handover_global_service_governance_zoho_desk.docx"

PENDENTE = "[PENDENTE DE CONFIRMAÇÃO]"

COLORS = {
    "blue": RGBColor(46, 116, 181),
    "dark_blue": RGBColor(31, 77, 120),
    "navy": RGBColor(11, 37, 69),
    "muted": RGBColor(89, 89, 89),
    "light_gray": "F2F4F7",
    "blue_gray": "E8EEF5",
    "callout": "F4F6F9",
    "white": "FFFFFF",
    "border": "C9D1D9",
}

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def paragraph_bottom_border(paragraph, color="2E74B5", size="8", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="C9D1D9", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = False

    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell, **CELL_MARGINS_DXA)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, before=0, after=0, line=1.05)

    set_table_borders(table)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_cell_text(cell, text, bold=False, color=None, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    set_paragraph_spacing(p, before=0, after=0, line=1.05)
    for i, part in enumerate(str(text).split("\n")):
        if i:
            p.add_run().add_break()
        run = p.add_run(part)
        set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        shade_cell(cell, COLORS["light_gray"])
        set_table_cell_text(cell, header, bold=True, color=COLORS["navy"], size=9.2)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_table_cell_text(cells[idx], value, size=font_size)

    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    set_paragraph_spacing(after, before=0, after=4, line=1.0)
    return table


def add_callout(doc, title, body, fill=COLORS["callout"]):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=2, line=1.1)
    run = p.add_run(title)
    set_run_font(run, size=10.5, color=COLORS["navy"], bold=True)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, before=0, after=0, line=1.1)
    run2 = p2.add_run(body)
    set_run_font(run2, size=10.2, color=RGBColor(35, 35, 35))
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=3, line=1.0)


def add_para(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    set_paragraph_spacing(p, before=0, after=6, line=1.1)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, bold=True)
        rest = text[len(bold_prefix):]
        if rest:
            p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, before=0, after=4, line=1.167)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(p, before=0, after=4, line=1.167)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run(item)


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    return p


def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    return p


def add_page_break(doc):
    doc.add_page_break()


def add_field(paragraph, instr):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(text)
    run._r.append(fld_char_end)


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    h1 = styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.color.rgb = COLORS["blue"]
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.line_spacing = 1.1

    h2 = styles["Heading 2"]
    h2.font.size = Pt(13)
    h2.font.color.rgb = COLORS["blue"]
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.1

    h3 = styles["Heading 3"]
    h3.font.size = Pt(12)
    h3.font.color.rgb = COLORS["dark_blue"]
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing = 1.1

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def configure_section(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    run = p.add_run("Global Service Governance - ZohoDesk")
    set_run_font(run, size=9, color=COLORS["muted"], bold=True)
    paragraph_bottom_border(p, color="D7DBE2", size="4", space="4")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Página ")
    set_run_font(run, size=9, color=COLORS["muted"])
    add_field(p, "PAGE")


def build_cover(doc):
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=18, after=18, line=1.0)

    kicker = doc.add_paragraph()
    set_paragraph_spacing(kicker, before=0, after=8, line=1.0)
    run = kicker.add_run("DOCUMENTAÇÃO FINAL DE PROJETO")
    set_run_font(run, size=10, color=COLORS["blue"], bold=True)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=0, after=4, line=1.0)
    run = title.add_run("Global Service Governance - ZohoDesk")
    set_run_font(run, size=26, color=COLORS["navy"], bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, before=0, after=16, line=1.1)
    run = subtitle.add_run(
        "Padronização Global de Atendimento, Governança Operacional e Evolução da Gestão de CX/Service"
    )
    set_run_font(run, size=13.5, color=COLORS["muted"], italic=True)

    rule = doc.add_paragraph()
    set_paragraph_spacing(rule, before=0, after=12, line=1.0)
    paragraph_bottom_border(rule, color="2E74B5", size="10", space="8")

    metadata = [
        ("Elaborado por", "Guilherme Bernardes Alamino"),
        ("Área", "CX / Pós-Vendas / Service"),
        ("Empresa", "Alliage"),
        ("Período do projeto", "Março/2026 a Julho/2026"),
        ("Versão", "1.0"),
        ("Data de emissão", "18 de junho de 2026"),
        ("Status do documento", f"Encerramento e handover em preparação; data final de aceite: {PENDENTE}"),
    ]
    add_table(doc, ["Campo", "Informação"], metadata, [2300, 7060], font_size=10)

    add_callout(
        doc,
        "Finalidade do documento",
        "Consolidar o que foi entregue, documentar o modelo operacional implantado e orientar a transição das rotinas para os responsáveis pós-projeto.",
        fill=COLORS["blue_gray"],
    )

    add_page_break(doc)


def build_document_map(doc):
    add_h1(doc, "Mapa do Documento")
    add_para(
        doc,
        "Este handover foi estruturado para servir simultaneamente como registro executivo, guia operacional e base de sustentação pós-projeto.",
    )
    rows = [
        ("1", "Sumário executivo", "Visão sintética do problema, solução e valor gerado."),
        ("2-7", "Contexto, objetivos, escopo e metodologia", "Base de entendimento do projeto e das ondas regionais."),
        ("8-10", "Entregas, processos e indicadores", "Detalhamento do modelo implementado e das métricas de gestão."),
        ("11-13", "Playbook, automações, IA e Knowledge Base", "Sustentação tecnológica e evolução do modelo."),
        ("14-20", "Riscos, aprendizados, handover e recomendações", "Continuidade operacional, backlog e próximos passos."),
        ("Anexos", "Lista sugerida", "Evidências e materiais complementares recomendados."),
    ]
    add_table(doc, ["Bloco", "Conteúdo", "Finalidade"], rows, [900, 3050, 5410])


def build_summary(doc):
    add_h1(doc, "1. Sumário Executivo")
    add_para(
        doc,
        "O projeto Global Service Governance - ZohoDesk foi conduzido para criar um modelo mínimo global de governança de atendimento, usando o ZohoDesk como plataforma central e o Playbook Global como camada de padronização, treinamento e sustentação.",
    )
    add_para(
        doc,
        "Antes do projeto, a operação global apresentava maturidade desigual entre regiões, uso não padronizado de status e campos, baixa comparabilidade de indicadores e dependência de controles paralelos. Isso dificultava a leitura de SLA, backlog, aging, qualidade de dados e priorização operacional.",
    )
    add_para(
        doc,
        "A solução construída consolidou uma arquitetura operacional composta por fluxo global, kanban/status oficiais, matriz de prioridade P1-P5, campos obrigatórios, governança por rituais, dashboard KPIv2, automação de Daily Backlog e documentação de suporte no Playbook Global.",
    )
    add_para(
        doc,
        "As regiões consideradas no modelo e nos relatórios são Brasil, Argentina, México, LATAM, USA e ROW. A implantação foi planejada por ondas para permitir diagnóstico, ajustes locais, treinamento, acompanhamento inicial e transição progressiva.",
    )
    add_callout(
        doc,
        "Valor gerado para a Alliage",
        "O projeto criou uma base comum para comparar performance regional, reduzir variação operacional, melhorar qualidade dos dados, estruturar rotinas de backlog e preparar a operação para automações, IA e gestão orientada a indicadores.",
    )
    add_h2(doc, "1.1 Principais entregas consolidadas")
    add_bullets(
        doc,
        [
            "Modelo global de status, fluxo e transições para tickets de atendimento.",
            "Matriz de prioridade por solicitante, tipo de atendimento, categoria e produto.",
            "Matriz consolidada de campos obrigatórios, recomendados, automáticos e condicionais.",
            "Dashboard KPIv2 com leituras Daily, Weekly, Monthly e Sync Health.",
            "Documentação operacional do Playbook Global em módulos de KPI, Kanban, Fluxo, Prioridade, Campos, Governança, Canais e Zoho Desk.",
            "Automação Daily Backlog Global com geração de PDFs globais/regionais e rascunho no Outlook.",
            "Base inicial para IA, Knowledge Base e ZIA Agents, com limites operacionais definidos.",
            "Backlog pós-projeto, riscos residuais e matriz de responsabilidades para sustentação.",
        ],
    )
    add_h2(doc, "1.2 Pontos pendentes críticos para aceite final")
    rows = [
        ("Data formal de encerramento", "Confirmar data de aceite/handover final em Julho/2026.", PENDENTE),
        ("Donos nominais pós-projeto", "Definir nomes por rotina, região e backup operacional.", PENDENTE),
        ("Evidências finais", "Anexar prints, links, atas, gravações e materiais de treinamento.", PENDENTE),
        ("Status produtivo da IA/ZIA", "Confirmar se ZIA Agents e primeira resposta SAC estão em produção, piloto ou desenho.", PENDENTE),
        ("Conta corporativa padrão", "Definir ownership de Supabase, Playbook, automações e credenciais.", PENDENTE),
    ]
    add_table(doc, ["Pendência", "Descrição", "Status"], rows, [2600, 4860, 1900])
    add_page_break(doc)


def build_context_objectives_scope(doc):
    add_h1(doc, "2. Contexto Inicial do Projeto")
    add_para(
        doc,
        "O cenário inicial combinava operação global de atendimento, diferentes níveis de maturidade regional e baixa padronização do ZohoDesk. Cada país possuía práticas próprias de registro, status, campos, rotinas de acompanhamento e uso de controles locais.",
    )
    add_bullets(
        doc,
        [
            "Uso não padronizado do ZohoDesk e baixa rastreabilidade entre abertura, atendimento e fechamento.",
            "Status diferentes entre países e ausência de critérios únicos para filas de espera.",
            "Campos obrigatórios inconsistentes, dificultando roteamento, prioridade, SLA e auditoria.",
            "Ausência de criticidade padronizada para ordenar filas e priorizar tickets de maior impacto.",
            "Dificuldade para acompanhar SLA, backlog, aging, tickets sem dono e tickets sem primeira resposta.",
            "Dependência de planilhas locais e controles paralelos para condução da rotina.",
            "Baixa comparabilidade entre regiões para análise executiva e tomada de decisão.",
        ],
    )

    add_h1(doc, "3. Objetivos do Projeto")
    add_h2(doc, "3.1 Objetivo geral")
    add_para(
        doc,
        "Implementar um modelo global mínimo de governança de atendimento utilizando o ZohoDesk como plataforma central, garantindo maior padronização, controle operacional, qualidade de dados e comparabilidade dos indicadores entre regiões.",
    )
    add_h2(doc, "3.2 Objetivos específicos")
    add_bullets(
        doc,
        [
            "Padronizar o fluxo global de atendimento e as transições de status.",
            "Definir campos obrigatórios, recomendados, automáticos e condicionais.",
            "Criar uma matriz de prioridade/criticidade objetiva e replicável.",
            "Definir leituras de SLA, backlog, aging, MTFC, MTTS e qualidade de dados.",
            "Estruturar dashboard global com recortes Daily, Weekly e Monthly.",
            "Criar rotina de governança operacional com rituais e donos claros.",
            "Implementar o projeto por ondas regionais, com treinamento e acompanhamento.",
            "Criar Playbook Global como referência única de processo, operação e sustentação.",
            "Preparar base para automações, IA, Knowledge Base e evolução tecnológica.",
            "Estruturar handover pós-projeto com responsáveis, riscos residuais e backlog.",
        ],
    )

    add_h1(doc, "4. Escopo e Não Escopo")
    rows_scope = [
        ("Padronização global do uso do ZohoDesk", "Incluído", "Fluxo, status, campos, prioridade, SLA e governança."),
        ("Campos obrigatórios", "Incluído", "Matriz por ticket e contato/cliente, com regras condicionais."),
        ("Status padronizados", "Incluído", "Modelo oficial de kanban, incluindo esperas específicas e fechamento."),
        ("KPIs e dashboard", "Incluído", "Daily, Weekly, Monthly, backlog, aging, MTFC, MTTS, SLA e sync health."),
        ("Playbook Global", "Incluído", "Módulos de operação, governança, canais, campos, prioridade e Zoho Desk."),
        ("Treinamentos regionais", "Incluído", f"Datas e evidências finais: {PENDENTE}."),
        ("IA e automações", "Incluído como base inicial", "Daily Backlog implementado; ZIA/KB com detalhes finais pendentes."),
        ("Operação global 24/7", "Não escopo", "Não foi desenhada como obrigação nesta fase."),
        ("Reestruturação organizacional", "Não escopo", "O projeto define governança e papéis, não organograma."),
        ("Integração completa ZohoDesk-SAP", "Não escopo", "Pode compor roadmap futuro."),
        ("Uso obrigatório por distribuidores", "Não escopo", "Nesta fase, foco no mínimo viável global."),
        ("IA resolutiva completa", "Não escopo", "IA não substitui análise humana em casos críticos ou sensíveis."),
    ]
    add_table(doc, ["Item", "Classificação", "Observação"], rows_scope, [2950, 1800, 4610], font_size=8.8)


def build_method_timeline(doc):
    add_h1(doc, "5. Metodologia de Implantação por Ondas")
    add_para(
        doc,
        "A implantação foi estruturada por ondas regionais para reduzir risco, respeitar maturidades locais e permitir estabilização progressiva. Cada onda seguiu diagnóstico, mapeamento de gaps, alinhamento com responsáveis, ajustes no ZohoDesk, treinamento, acompanhamento e handover.",
    )
    add_numbers(doc, ["Brasil", "USA", "ROW", "Argentina / LATAM", "México"])

    rows = [
        ("Diagnóstico inicial", "Levantar operação atual, status, campos, canais, backlog e controles paralelos.", "Lista de gaps por região."),
        ("Mapeamento de gaps", "Comparar prática local contra o modelo mínimo global.", "Plano de ajuste regional."),
        ("Alinhamento local", "Validar owners, rituais e adaptações permitidas.", "Responsáveis e agenda de implantação."),
        ("Configuração/ajustes", "Aplicar campos, status, regras, dashboard e documentação de apoio.", "Ambiente pronto para treinamento."),
        ("Treinamento", "Capacitar times regionais em fluxo, status, prioridade, campos e dashboard.", f"Evidências: {PENDENTE}."),
        ("Acompanhamento inicial", "Monitorar backlog, dúvidas, uso incorreto e qualidade de dados.", "Correções pós-go-live."),
        ("Handover", "Transferir rotina, riscos, backlog e ownership.", f"Aceite por responsável: {PENDENTE}."),
    ]
    add_table(doc, ["Etapa", "Atividade", "Saída esperada"], rows, [2200, 4050, 3110], font_size=8.8)

    add_h1(doc, "6. Linha do Tempo do Projeto")
    timeline = [
        ("Abertura e diagnóstico", "Março/2026", "Levantamento do cenário inicial e principais gaps.", "Problemas e oportunidades mapeados."),
        ("Modelo global", "Abril/2026", "Definição de campos, status, SLA, prioridade e governança.", "Estrutura global mínima definida."),
        ("Expansão regional", "Maio/2026", "Implementação USA/ROW e consolidação dos KPIs.", "Primeiras waves internacionais aplicadas."),
        ("LATAM/MEX, IA e automações", "Junho/2026", "Treinamentos, ajustes regionais e evolução tecnológica.", "Base para sustentação e escala."),
        ("Handover e encerramento", "Julho/2026", "Documentação final, responsáveis e backlog pós-projeto.", f"Projeto preparado para transição; aceite final {PENDENTE}."),
    ]
    add_table(doc, ["Fase", "Período", "Principais atividades", "Resultado"], timeline, [1900, 1400, 3400, 2660], font_size=8.5)
    add_page_break(doc)


def build_deliverables(doc):
    add_h1(doc, "7. Entregas Realizadas")
    add_para(
        doc,
        "A tabela abaixo consolida as entregas do projeto com o problema endereçado, resultado gerado, evidência esperada e responsável futuro. Quando a evidência ou o dono nominal não aparece confirmado no material disponível, o campo permanece pendente.",
    )
    rows = [
        ("1", "Padronização de campos obrigatórios", "Campos inconsistentes e baixa qualidade de dados.", "Matriz consolidada de ticket e contato/cliente.", "05_Campos_Obrigatorios", PENDENTE),
        ("2", "Padronização de status", "Uso genérico de status e filas pouco rastreáveis.", "Modelo oficial com Aberto, Em Atendimento, esperas, Resolvido e Fechado.", "02_Kanban", PENDENTE),
        ("3", "Definição de SLA e prioridade", "Priorização subjetiva e baixa comparabilidade.", "Matriz P1-P5, metas MTFC/MTTS e leitura de SLA.", "04_Prioridade / KPIv2", PENDENTE),
        ("4", "Dashboard global KPIv2", "Indicadores regionais não comparáveis.", "Abas Daily, Weekly, Monthly e Sync Health.", "01_KPI/KPI_V2", PENDENTE),
        ("5", "Global Playbook", "Dependência de conhecimento informal.", "Plataforma de referência multilíngue por módulos.", "README e módulos do playbook", PENDENTE),
        ("6", "Governança operacional", "Ausência de cadência e owner para análise.", "Daily, quinzenal, mensal e quarterly.", "06_Governanca", PENDENTE),
        ("7", "Rotina de backlog", "Backlog sem leitura comum e baixa visibilidade de risco.", "Daily Backlog Global e PDFs por região.", "reports/ e docs/02-operacao/daily-backlog", PENDENTE),
        ("8", "Treinamentos regionais", "Adoção desigual do modelo.", "Materiais e evidências de treinamento.", f"{PENDENTE}", PENDENTE),
        ("9", "Auditoria de dados", "Campos e status sem validação prática.", "Método de amostragem e plano de melhoria.", "06_Governanca/auditoria", PENDENTE),
        ("10", "Base para IA / Knowledge Base", "Atendimento inicial sem base estruturada.", "Arquitetura e limites definidos.", f"{PENDENTE}", PENDENTE),
        ("11", "ZIA Agents para primeira resposta SAC", "Necessidade de escala em respostas elegíveis.", "Desenho operacional e restrições de segurança.", f"{PENDENTE}", PENDENTE),
        ("12", "Automações de reports", "Montagem manual de relatórios.", "Daily Backlog automatizado; weekly/monthly em roadmap.", "tools/run-daily-backlog-report.js", PENDENTE),
        ("13", "Documentação operacional", "Baixa rastreabilidade do processo.", "DOCX, Markdown, páginas HTML e documentação técnica.", "docs/ e módulos", PENDENTE),
        ("14", "Processo de handover", "Risco de dependência do trainee.", "Documento final, matriz de responsabilidades e backlog.", "Este documento", PENDENTE),
    ]
    add_table(
        doc,
        ["ID", "Entrega", "Problema que resolveu", "Resultado gerado", "Evidência/documento", "Responsável futuro"],
        rows,
        [450, 1750, 2200, 2360, 1600, 1000],
        font_size=7.6,
    )


def build_processes(doc):
    add_h1(doc, "8. Processos Definidos")
    add_h2(doc, "8.1 Fluxo Global de Atendimento")
    add_para(
        doc,
        "O fluxo global organiza o ciclo do ticket desde a entrada até o encerramento, mantendo rastreabilidade entre triagem, atendimento, esperas, resolução e fechamento. O objetivo é que o status do ticket represente a condição real do atendimento, sem uso genérico de filas.",
    )
    rows = [
        ("Entrada do ticket", "Ticket nasce por canal oficial ou registro interno.", "Dados mínimos preenchidos e ticket pronto para triagem."),
        ("Triagem", "Validação de solicitante, tipo, categoria, produto e região.", "Direcionamento correto e prioridade calculável."),
        ("Classificação", "Aplicação de campos obrigatórios e matriz de prioridade.", "Ticket com contexto suficiente para fila e SLA."),
        ("Atendimento", "Agente ou especialista executa ação técnica ou SAC.", "Atualizações frequentes e plano de ação claro."),
        ("Tratamento de pendências", "Uso de status de espera específico: cliente, peça ou terceiro/visita.", "Causa de bloqueio visível e auditável."),
        ("Resolução", "Registro da solução aplicada e evidências relevantes.", "Caso pronto para validação de fechamento."),
        ("Fechamento", "Encerramento definitivo com resumo de resolução.", "Ticket concluído, rastreável e apto para KPI/auditoria."),
        ("Reabertura", "Aplicável quando há retrabalho ou pendência após fechamento.", f"Critério e motivo de reabertura: {PENDENTE}."),
    ]
    add_table(doc, ["Etapa", "Descrição", "Resultado esperado"], rows, [2100, 3850, 3410], font_size=8.5)

    add_h2(doc, "8.2 Status Padronizados")
    rows = [
        ("Aberto", "Na abertura, com informações mínimas para iniciar atendimento.", "Permanência curta até primeiro tratamento.", "Backlog oculto de entrada se ficar parado."),
        ("Em Atendimento", "Quando há ação técnica efetiva em andamento.", "Atualizações frequentes e avanço concreto.", "Mascara filas de espera se usado como genérico."),
        ("Aguardando Cliente", "Quando o próximo passo depende do cliente.", "Motivo de espera e contexto registrados.", "Aging sem controle se não houver follow-up."),
        ("Aguardando Peça", "Quando há dependência de peça, material ou logística.", "Pedido rastreável e previsão acompanhada.", "Gargalo logístico invisível se não monitorado."),
        ("Aguardando AT/Terceiro", "Quando depende de assistência técnica, terceiro ou visita.", "Dependência e previsão registradas.", "Bloqueio externo sem dono se mal usado."),
        ("Resolvido", "Quando a tratativa terminou e aguarda validação de encerramento.", "Resumo da solução e evidências.", "Distorce qualidade se for pulado."),
        ("Fechado", "Após resolução e validação das condições de encerramento.", "Ticket concluído e auditável.", "Fechamento precoce gera retrabalho/reabertura."),
    ]
    add_table(doc, ["Status", "Quando usar", "Comportamento esperado", "Impacto em SLA/backlog"], rows, [1650, 2800, 2550, 2360], font_size=8.0)
    add_callout(
        doc,
        "Regra de transição",
        "As transições válidas preservam rastreabilidade: Aberto -> Em Atendimento; Em Atendimento -> esperas específicas ou Resolvido; esperas -> Em Atendimento; Resolvido -> Fechado.",
    )

    add_h2(doc, "8.3 Campos Obrigatórios")
    add_para(
        doc,
        "A matriz de campos separa dados de ticket e contato/cliente em quatro classificações: obrigatório, desejável, automático e condicional. Essa estrutura sustenta prioridade, fluxo, SLA, KPI e auditoria.",
    )
    rows = [
        ("Campos do ticket", "Nome do solicitante, e-mail, telefone, perfil do solicitante, tipo de atendimento, categoria, produto, marca, número de série, assunto e descrição.", "Abertura, classificação, prioridade e roteamento."),
        ("Campos do contato/cliente", "Nome, sobrenome, conta, e-mail, telefone, razão social, endereço, idioma, dealer e dados técnicos do equipamento.", "Qualidade cadastral, contato e rastreabilidade."),
        ("Campos automáticos", "ID do ticket, data de abertura, canal, owner, status, prioridade, SLA de primeira resposta, SLA de resolução, primeira resposta e fechamento.", "Medição de SLA, auditoria e dashboards."),
        ("Campos condicionais", "Estado para Brasil, província para Argentina, nome da assistência/distribuidor, número de série com exceção contextual.", "Adaptação regional sem perder padrão global."),
        ("Campos de fechamento", "Resumo de resolução, motivo de fechamento, solução remota/presencial, peça utilizada e quantidade final de peças.", "Qualidade de encerramento, auditoria e lições operacionais."),
    ]
    add_table(doc, ["Grupo", "Campos/regras principais", "Uso operacional"], rows, [2000, 4700, 2660], font_size=8.3)

    add_h2(doc, "8.4 Priorização e Criticidade")
    add_para(
        doc,
        "A prioridade foi estruturada como cálculo objetivo baseado em quatro eixos: solicitante, tipo de atendimento, categoria e produto. A pontuação final é a multiplicação dos pesos e gera uma faixa P1-P5.",
    )
    rows = [
        ("Solicitante", "Assistência/Distribuidor = 1; Cliente = 2; Cliente de Locação = 3."),
        ("Tipo de atendimento", "Dentista especialista de produto = 1; SAC = 2; Técnico especializado = 3; Cliente de locação com suporte técnico = 4."),
        ("Categoria", "Dúvidas gerais = 1; Instalação linha de imagem = 2; Dúvidas sobre equipamento = 3; Assuntos financeiros = 4; Problemas técnicos = 5."),
        ("Produto", "P1 a P5 = peso 1; P6 a P10 = peso 2; P11 a P15 = peso 3."),
        ("Fórmula", "Pontuação = Solicitante x Tipo de Atendimento x Categoria x Produto."),
    ]
    add_table(doc, ["Eixo", "Regra"], rows, [2200, 7160], font_size=8.5)
    rows = [
        ("P1 - Crítica", "141-180", "Atendimento imediato e visibilidade de liderança."),
        ("P2 - Alta", "101-140", "Alta precedência e follow-up próximo."),
        ("P3 - Moderada", "61-100", "Tratativa priorizada dentro da rotina."),
        ("P4 - Baixa", "26-60", "Fluxo padrão, menor precedência."),
        ("P5 - Planejada", "1-25", "Baixa urgência relativa e tratamento planejado."),
    ]
    add_table(doc, ["Prioridade", "Faixa", "Leitura operacional"], rows, [1900, 1200, 6260], font_size=8.5)

    add_h2(doc, "8.5 SLA, Backlog, Aging e Governança")
    rows = [
        ("MTFC", "Tempo médio até primeira resposta.", "Meta por prioridade: Urgente 1h, Alta 2h, Média 3h, Baixa 5h, Muito Baixa 6h."),
        ("MTTS", "Tempo médio até solução/fechamento.", "Meta por região: Brasil 4d, Argentina 5d, México/LATAM 6d, USA/ROW 10d."),
        ("SLA compliance", "Percentual de tickets elegíveis dentro das metas.", "Usado em Weekly e Monthly para leitura de performance."),
        ("Backlog", "Tickets abertos no recorte operacional.", "Acompanhamento diário por região, status, prioridade e grupo."),
        ("Aging", "Idade dos tickets abertos.", "Identifica fila crítica, tickets antigos e risco de escalonamento."),
        ("Tickets sem dono", "Tickets sem agente ou com grupo operacional Sem dono.", "Base para cobrança de ownership e ajuste de distribuição."),
    ]
    add_table(doc, ["Indicador/processo", "Definição", "Regra/uso"], rows, [2100, 2900, 4360], font_size=8.3)

    add_h2(doc, "8.6 Rituais de Governança Operacional")
    rows = [
        ("Daily operacional", "Diário", "Supervisor local", "Fila priorizada, pendências com owner e escalonamentos imediatos."),
        ("Governança quinzenal", "Quinzenal", "Gestão regional", "Gargalos recorrentes, ajustes táticos e plano interáreas."),
        ("Governança mensal", "Mensal", "Liderança global de service", "Performance consolidada, melhorias estruturais e direcionamento regional."),
        ("Governança quarterly", "Trimestral", "Sponsor e liderança executiva", "Decisão estratégica, priorização institucional e remoção de bloqueios."),
        ("Auditoria de dados", "Ciclo definido por governança", "Global/Regional Owner", "Score de aderência, desvios críticos e plano de melhoria."),
    ]
    add_table(doc, ["Ritual", "Frequência", "Responsável", "Saída esperada"], rows, [1900, 1300, 2250, 3910], font_size=8.3)

    add_h2(doc, "8.7 NPS / CSAT")
    add_para(
        doc,
        f"NPS e CSAT devem ser analisados junto com indicadores operacionais. A leitura recomendada é cruzar satisfação com SLA, MTFC, MTTS, backlog, reabertura e qualidade de fechamento. Fonte, meta e periodicidade oficial de NPS/CSAT: {PENDENTE}.",
    )

    add_h2(doc, "8.8 Automação e IA no Processo")
    add_para(
        doc,
        "Automações e IA foram tratadas como evolução do modelo, não como substituição da governança. A premissa é que automação amplifica um processo já padronizado; se os dados de entrada estiverem incorretos, dashboards, relatórios e IA também serão impactados.",
    )
    add_page_break(doc)


def build_dashboards_playbook_automation_ai(doc):
    add_h1(doc, "9. Dashboards e Indicadores")
    add_para(
        doc,
        "O dashboard KPIv2 consolida leituras Daily, Weekly, Monthly e Sync Health. A fonte oficial de dados é o Zoho Desk, sincronizado para Supabase e consumido pela Edge Function dashboard-read. O frontend aplica filtros, cache e regras de negócio para exibir os indicadores operacionais.",
    )
    rows = [
        ("Zoho Desk", "Sistema fonte", "Tickets, detalhes, métricas, departamentos, agentes, contatos e histórico de status."),
        ("sync-tickets-v0", "Ingestão", "Busca dados via API Zoho e grava no Supabase."),
        ("Supabase public.zoho_*", "Armazenamento", "Tabelas principais de tickets, agentes, departamentos, contatos, status history e sync."),
        ("vw_tickets_bi_base", "Camada BI canônica", "Normaliza região, prioridade, grupo, MTFC, MTTS, aging, metas e status de SLA."),
        ("dashboard-read", "API de leitura", "Endpoint para views bi-tickets, bi-backlog, bi-kpis e sync-health."),
        ("KPIv2", "Frontend", "Daily, Weekly, Monthly, filtros, gráficos, tabelas e exportações."),
    ]
    add_table(doc, ["Camada", "Papel", "Descrição"], rows, [2100, 1800, 5460], font_size=8.3)

    add_h2(doc, "9.1 KPIs definidos")
    rows = [
        ("Volume de tickets", "Quantidade de tickets criados/fechados no período.", "Medir demanda e capacidade.", "Weekly/Monthly", "Global Service Owner / Regional Manager"),
        ("Tickets abertos", "Tickets ainda sem fechamento.", "Acompanhar carga operacional.", "Daily", "Supervisor local"),
        ("Tickets fechados", "Tickets com fechamento no período.", "Avaliar produção e vazão.", "Weekly/Monthly", "Regional Manager"),
        ("Backlog", "Total de tickets abertos dentro do escopo BI.", "Controlar fila pendente.", "Diário", "Supervisor local"),
        ("Aging", "Idade dos tickets abertos.", "Identificar tickets antigos e risco.", "Diário", "Supervisor local"),
        ("MTFC", "Tempo médio até primeira resposta.", "Medir velocidade de entrada no atendimento.", "Weekly/Monthly", "Regional Manager"),
        ("MTTS", "Tempo médio até solução.", "Medir tempo total de resolução.", "Weekly/Monthly", "Regional Manager"),
        ("SLA compliance", "Percentual dentro das metas.", "Controlar cumprimento de compromisso.", "Weekly/Monthly", "Global Service Owner"),
        ("CSAT", "Satisfação por atendimento.", "Avaliar qualidade percebida.", f"{PENDENTE}", f"{PENDENTE}"),
        ("NPS", "Lealdade/recomendação do cliente.", "Avaliar saúde relacional.", f"{PENDENTE}", f"{PENDENTE}"),
        ("Reabertura", "Tickets reabertos após fechamento.", "Medir retrabalho e qualidade de solução.", f"{PENDENTE}", "Regional Manager"),
        ("Tickets sem dono", "Tickets sem agente/grupo responsável.", "Corrigir distribuição e ownership.", "Diário", "Supervisor local"),
        ("Tickets em risco", "P1/P2, >10 dias ou sem primeira resposta.", "Priorização operacional imediata.", "Diário", "Supervisor local"),
    ]
    add_table(doc, ["KPI", "Definição", "Objetivo", "Frequência", "Responsável"], rows, [1500, 2600, 2350, 1450, 1460], font_size=7.4)

    add_h2(doc, "9.2 Observações de consistência")
    add_bullets(
        doc,
        [
            "Daily/Backlog usa backlog aberto atual; não usa semana, mês ou intervalo de datas.",
            "Weekly calcula criados por created_time e fechados/SLA/MTFC/MTTS por closed_time na semana.",
            "Monthly usa mês/ano como período oficial para tickets fechados; os campos De/Até não alteram os números atuais.",
            "O card Backlog do Weekly vem do backlog canônico atual e não muda com a semana selecionada.",
            "A base BI inclui grupos operacionais Suporte geral, Especialista e Sem dono, excluindo grupos fora do escopo de reporte.",
            "A automação e o dashboard devem ser reconciliados para evitar divergências entre report e tela.",
        ],
    )

    add_h1(doc, "10. Playbook Global")
    add_para(
        doc,
        "O Playbook Global é a referência central de sustentação do modelo. Ele organiza o conhecimento operacional em módulos navegáveis e multilíngues, reduzindo dependência de conhecimento informal e apoiando treinamento, auditoria e replicação do modelo.",
    )
    rows = [
        ("KPI", "Daily, Weekly, Monthly, documentação de números, filtros e linhagem de dados."),
        ("Kanban", "Status oficiais, objetivos, transições válidas e sinais de gestão."),
        ("Fluxo Global", "Etapas de atendimento, campos por etapa e adaptações locais."),
        ("Prioridade", "Matriz P1-P5, pesos, fórmula e simulação."),
        ("Campos Obrigatórios", "Matriz consolidada de ticket e contato/cliente."),
        ("Governança", "Rituais, responsabilidades, auditoria e indicadores de maturidade."),
        ("Canais", "Telefone, WhatsApp, formulário e e-mail, com garantia de campos mínimos."),
        ("Zoho Desk", "Tutorial, edição/administração e conteúdo legado preservado."),
    ]
    add_table(doc, ["Módulo", "Conteúdo principal"], rows, [2200, 7160], font_size=8.5)

    add_h1(doc, "11. Automações Desenvolvidas ou Planejadas")
    add_para(
        doc,
        "A automação mais consolidada é o Daily Backlog Global PDF. Ela executa o utilitário local, consulta a fonte de dados do dashboard, gera PDFs global e regionais, cria manifesto, corpo trilingue e payload de Outlook, e prepara rascunho sem envio automático.",
    )
    rows = [
        ("Daily Backlog Global PDF", "Padronizar leitura diária do backlog global e regional.", "Diária", "dashboard-read / Supabase / KPIv2", "Lista padrão do report", "Guilherme até handover; futuro dono pendente", "Implementada com handover documentado"),
        ("Weekly report", "Consolidar performance semanal por região, SLA, MTFC, MTTS e backlog.", "Semanal", "KPIv2 Weekly / Supabase", PENDENTE, PENDENTE, "Planejada/em validação"),
        ("Monthly report global/regional", "Leitura executiva mensal de performance e tendência.", "Mensal", "KPIv2 Monthly / Supabase", PENDENTE, PENDENTE, "Planejada/em validação"),
        ("Fallback mensal", "Rodar no primeiro dia útil seguinte se a automação mensal falhar.", "Mensal sob exceção", "Automação mensal", PENDENTE, PENDENTE, "Planejada"),
        ("Reconciliação report x dashboard", "Garantir consistência entre números enviados e números exibidos.", "Por ciclo de report", "Manifest.json, dashboard e Supabase", "Governança", PENDENTE, "Recomendação pós-projeto"),
    ]
    add_table(
        doc,
        ["Automação", "Objetivo", "Frequência", "Fonte de dados", "Destinatários", "Responsável futuro", "Status"],
        rows,
        [1350, 2100, 950, 1550, 1200, 1250, 960],
        font_size=7.0,
    )
    add_callout(
        doc,
        "Critério de sucesso da automação Daily Backlog",
        "A execução só deve ser considerada concluída quando a pasta oficial existir, os PDFs e arquivos textuais forem gerados, o rascunho no Outlook existir e todos os anexos esperados forem conferidos no draft salvo.",
    )

    add_h1(doc, "12. IA e Knowledge Base")
    add_para(
        doc,
        "A frente de IA e Knowledge Base deve ser entendida como etapa de evolução do projeto. O modelo recomendado parte de base de conhecimento validada, classificação de intenção, primeira resposta apenas para casos elegíveis e encaminhamento humano quando houver risco, exceção ou sensibilidade operacional.",
    )
    rows = [
        ("Knowledge Base", "Artigos por intenção, idioma, produto, categoria e cenário de atendimento.", f"Conteúdo final e owner editorial: {PENDENTE}."),
        ("Primeira resposta SAC", "Resposta inicial padronizada para demandas elegíveis e de baixo risco.", f"Status produtivo/piloto: {PENDENTE}."),
        ("Classificação de tickets", "Identificação de intenção, categoria e encaminhamento sugerido.", f"Regras finais: {PENDENTE}."),
        ("Casos elegíveis", "Dúvidas gerais, orientações de processo, pedido de informação e suporte inicial de baixo risco.", f"Lista final: {PENDENTE}."),
        ("Casos não elegíveis", "Garantia, promessa de prazo, caso crítico, exceção comercial, reclamação sensível ou decisão técnica complexa.", "Sempre encaminhar para humano."),
        ("Regras de segurança", "IA não aprova/nega garantia, não promete prazo, não encerra caso crítico e não substitui análise humana.", "Obrigatório manter em qualquer desenho."),
        ("Linguagem e acentuação", "Respostas revisadas em PT-BR, EN e ES conforme operação regional.", f"Guia de tom final: {PENDENTE}."),
    ]
    add_table(doc, ["Componente", "Definição operacional", "Status/pendência"], rows, [2200, 4300, 2860], font_size=8.2)
    add_h2(doc, "12.1 Arquitetura resumida recomendada")
    add_numbers(
        doc,
        [
            "Ticket criado no ZohoDesk.",
            "Workflow aciona Deluge.",
            "Deluge chama ZIA Agent.",
            "ZIA consulta o ticket e a Knowledge Base.",
            "ZIA classifica a demanda e valida elegibilidade.",
            "ZIA envia primeira resposta quando elegível.",
            "Ticket segue para agente humano com contexto registrado.",
        ],
    )
    add_page_break(doc)


def build_risks_lessons_skills(doc):
    add_h1(doc, "13. Riscos Identificados e Mitigados")
    rows = [
        ("Resistência das regiões", "Baixa adesão ao padrão global.", "Implantação por ondas, adaptação local controlada e Playbook de referência.", "Aderência pode cair sem governança.", "Regional Manager / Global Service Owner"),
        ("Maturidade desigual", "Diferença de uso e disciplina entre países.", "Modelo mínimo global e rituais por nível.", "Necessário monitorar adoção.", "Global Service Owner"),
        ("Baixa qualidade dos dados", "KPIs distorcidos e SLA pouco confiável.", "Campos obrigatórios, auditoria e matriz consolidada.", "Risco permanece se campos forem burlados.", "Supervisor local"),
        ("Uso incorreto do ZohoDesk", "Status e filas deixam de refletir a realidade.", "Treinamento, status oficiais e auditoria por amostra.", "Requer reciclagem periódica.", "Supervisor local"),
        ("Falta de sustentação pós-projeto", "Retorno a controles paralelos.", "Matriz de responsabilidades, rituais e backlog pós-projeto.", "Donos nominais pendentes.", PENDENTE),
        ("Dependência do trainee", "Conhecimento concentrado em uma pessoa.", "Documentação operacional e handover formal.", "Necessário owner corporativo.", PENDENTE),
        ("Falhas em automações", "Report incompleto ou sem anexos.", "Validação de artefatos e anexos no Outlook.", "Manter monitoramento e fallback.", PENDENTE),
        ("IA responder fora do escopo", "Risco de orientação indevida ao cliente.", "Limites de segurança e encaminhamento humano.", "Depende da KB e governança de conteúdo.", PENDENTE),
        ("Divergência dashboard x ZohoDesk", "Perda de confiança nos números.", "Linhagem Zoho -> Supabase -> BI -> dashboard e sync health.", "Reconciliar periodicamente.", "Dono do dashboard"),
        ("Crescimento de escopo", "Atraso e perda de foco no mínimo viável.", "Escopo e não escopo documentados.", "Roadmap futuro precisa priorização.", "Sponsor"),
    ]
    add_table(doc, ["Risco", "Impacto", "Mitigação realizada", "Risco residual", "Dono futuro"], rows, [1700, 1800, 2750, 1900, 1210], font_size=7.4)
    add_callout(
        doc,
        "Quatro riscos principais para sustentação",
        "Compreensão desigual do treinamento, governança sem rotina, dados inconsistentes no BI/Supabase e uso incorreto do ZohoDesk devem permanecer como foco da rotina pós-projeto.",
    )

    add_h1(doc, "14. Lições Aprendidas")
    add_bullets(
        doc,
        [
            "Antes de automatizar, é necessário padronizar o processo.",
            "Dashboard só é confiável se o dado de entrada estiver correto.",
            "Treinamento precisa ser acompanhado por auditoria prática.",
            "O mínimo comum global é mais viável do que tentar impor todos os detalhes locais.",
            "A governança precisa ter dono, frequência e evidência.",
            "IA depende de base de conhecimento estruturada, revisada e com limites claros.",
            "Regiões diferentes exigem adaptação, mas sem perder o padrão global.",
            "Indicadores sem rotina de análise não geram mudança operacional.",
            "Handover precisa começar antes do final do projeto, com backlog e owners já definidos.",
            "Automação de report precisa validar não só geração de arquivo, mas também entrega e anexos.",
        ],
    )

    add_h1(doc, "15. Habilidades Desenvolvidas")
    rows = [
        ("Gestão de projetos", "Planejamento por ondas, priorização, gestão de riscos, handover e comunicação executiva."),
        ("Processos e governança", "Padronização global, definição de rituais, construção de playbook e auditoria de processo."),
        ("Dados e indicadores", "Definição de KPIs, dashboard, SLA, backlog, aging e comparabilidade global."),
        ("Sistemas e tecnologia", "ZohoDesk, Supabase, dashboard web, automação com Codex/Claude, Deluge, ZIA Agents e Knowledge Base."),
        ("Comunicação e stakeholders", "Condução com Brasil, USA, ROW, LATAM, Argentina e México; alinhamento com diretoria; treinamentos e gestão de resistências."),
    ]
    add_table(doc, ["Dimensão", "Habilidades evidenciadas"], rows, [2300, 7060], font_size=8.7)


def build_handover_backlog_recommendations(doc):
    add_h1(doc, "16. Responsáveis Pós-Projeto")
    add_para(
        doc,
        "A matriz abaixo define responsabilidades operacionais por atividade. Os papéis foram indicados com base no modelo de governança; nomes nominais, backups e aceite formal precisam ser confirmados no encerramento.",
    )
    rows = [
        ("Monitoramento do backlog", "Supervisor local", "Regional Manager", "Diária", "Lista priorizada e ações registradas."),
        ("Revisão semanal dos KPIs", "Regional Manager", "Global Service Owner", "Semanal", "Resumo semanal e plano de ação."),
        ("Atualização de campos/status no ZohoDesk", "Global Service Owner", "Administrador ZohoDesk", "Sob demanda controlada", "Registro de mudança e comunicação."),
        ("Auditoria de qualidade dos dados", "Global/Regional Owner", "Supervisor local", "Ciclo definido", "Score, desvios e plano de melhoria."),
        ("Manutenção do dashboard", "Dono do dashboard", "TI/Dados", "Contínua", "Checklist de sync health e validações."),
        ("Manutenção do playbook", "Global Service Owner", "Owners de módulo", "Mensal ou por mudança", "Versão atualizada e changelog."),
        ("Atualização da Knowledge Base", "Owner editorial KB", "Especialistas técnicos", "Contínua", "Artigos revisados e aprovados."),
        ("Monitoramento da ZIA", "Owner IA/ZIA", "Service Owner", "Diária/semana inicial", "Logs, exceções e taxa de encaminhamento humano."),
        ("Envio de reports", "Dono da automação", "Backup operacional", "Diária/Semanal/Mensal", "Draft/relatório gerado e conferido."),
        ("Treinamento de novos usuários", "Supervisor local", "Global Service Owner", "Onboarding", "Lista de presença e avaliação."),
        ("Planos de ação regionais", "Regional Manager", "Supervisor local", "Quinzenal/Mensal", "Plano com owner e prazo."),
        ("Gestão de NPS/CSAT", PENDENTE, "Regional Manager", PENDENTE, "Relatório de satisfação e ações."),
    ]
    add_table(doc, ["Atividade", "Responsável principal", "Apoio/backup", "Frequência", "Evidência esperada"], rows, [1850, 1800, 1700, 1450, 2560], font_size=7.4)

    add_h1(doc, "17. Backlog Pós-Projeto")
    rows = [
        ("B-001", "Finalizar ajustes BRA/ROW/LATAM/MEX.", "Operacional", "Alta", PENDENTE, "Julho/2026", "Validar lista final por região."),
        ("B-002", "Consolidar treinamento LATAM/MEX.", "Treinamento", "Alta", PENDENTE, "Julho/2026", "Anexar evidências e dúvidas recorrentes."),
        ("B-003", "Revisar gargalo do México.", "Processo", "Alta", PENDENTE, "Julho/2026", "Detalhar causa raiz e plano."),
        ("B-004", "Finalizar automações de reports weekly/monthly.", "Automação", "Alta", PENDENTE, "Agosto/2026", "Garantir consistência com dashboard."),
        ("B-005", "Validar dashboard com dados do ZohoDesk.", "Dados", "Alta", "Dono do dashboard", "Contínuo", "Comparar amostras e sync health."),
        ("B-006", "Documentar rotina de IA.", "IA/KB", "Média", PENDENTE, "Agosto/2026", "Incluir elegibilidade, exceções e logs."),
        ("B-007", "Definir conta padrão para Supabase/Playbook/automações.", "Governança técnica", "Alta", PENDENTE, "Julho/2026", "Reduzir risco de dependência individual."),
        ("B-008", "Criar plano de sustentação pós-handover.", "Governança", "Alta", "Global Service Owner", "Julho/2026", "Rotina de 30/60/90 dias."),
        ("B-009", "Revisar categorias futuras para encaminhamento comercial.", "Processo", "Média", PENDENTE, "Agosto/2026", "Evitar desvio de escopo no service."),
        ("B-010", "Evoluir base de conhecimento SAC/ST.", "IA/KB", "Média", PENDENTE, "Contínuo", "Priorização por volume e risco."),
    ]
    add_table(
        doc,
        ["ID", "Atividade pendente", "Tipo", "Prioridade", "Responsável", "Prazo sugerido", "Observação"],
        rows,
        [700, 2600, 1050, 1000, 1250, 1200, 1560],
        font_size=7.0,
    )

    add_h1(doc, "18. Recomendações Futuras")
    add_numbers(
        doc,
        [
            "Manter auditoria semanal de dados até estabilização da adoção por região.",
            "Criar dono formal para o dashboard e para a camada Supabase.",
            "Revisar mensalmente SLAs por região, considerando maturidade, volume e capacidade.",
            "Formalizar sustentação das automações com owner, backup e rotina de validação.",
            "Migrar Supabase, Playbook e automações para conta padrão corporativa.",
            "Evoluir IA apenas com Knowledge Base validada e política de elegibilidade aprovada.",
            "Criar rotina de reciclagem dos usuários e onboarding para novos agentes.",
            "Usar indicadores para planos de ação regionais, não apenas para reporte.",
            "Documentar qualquer alteração no ZohoDesk com owner, data, motivo e impacto esperado.",
            "Realizar revisão trimestral da governança global com sponsor e liderança executiva.",
        ],
    )

    add_h1(doc, "19. Conclusão")
    add_para(
        doc,
        "O projeto Global Service Governance - ZohoDesk deixou uma base estruturada de governança global para a operação de atendimento. A iniciativa aumentou a visibilidade da operação, criou comparabilidade entre regiões, reduziu dependência de controles paralelos e organizou um modelo mínimo comum para fluxo, status, prioridade, campos, indicadores e rituais.",
    )
    add_para(
        doc,
        "A sustentação do valor gerado dependerá de disciplina pós-projeto: donos claros, auditoria de dados, atualização do Playbook, monitoramento do dashboard, gestão ativa do backlog e governança das automações. A próxima etapa de maturidade deve priorizar ownership corporativo, reconciliação entre ZohoDesk/Supabase/dashboard e evolução segura de IA e Knowledge Base.",
    )
    add_callout(
        doc,
        "Mensagem final de handover",
        f"O projeto está preparado para transição operacional, desde que os responsáveis nominais, datas finais, evidências de treinamento, status da IA e conta corporativa padrão sejam confirmados: {PENDENTE}.",
        fill=COLORS["blue_gray"],
    )


def build_annexes(doc):
    add_h1(doc, "20. Anexos Sugeridos")
    rows = [
        ("A1", "TAP do projeto", "Formalizar objetivo, escopo, sponsors e critérios de sucesso."),
        ("A2", "Roadmap do projeto", "Evidenciar ondas, marcos, decisões e evolução temporal."),
        ("A3", "Global Playbook", "Referência operacional central para treinamento e auditoria."),
        ("A4", "Prints do dashboard", "Evidenciar KPIs Daily, Weekly, Monthly e Sync Health."),
        ("A5", "Evidências de treinamento", "Listas de presença, gravações, materiais e avaliações."),
        ("A6", "Documentação das automações", "Daily Backlog, reports, fallback e validação de anexos."),
        ("A7", "Documentação ZIA/Knowledge Base", "Arquitetura, artigos, elegibilidade, limites e logs."),
        ("A8", "Matriz de responsabilidades", "Responsáveis principais, backups e evidências esperadas."),
        ("A9", "Backlog pós-projeto", "Pendências, prioridades, prazos e owners."),
        ("A10", "Procedimentos relacionados", "SOPs, políticas internas e guias ZohoDesk."),
        ("A11", "Evidências de dados/Supabase", "Linhagem, views, endpoints, sync runs e reconciliações."),
        ("A12", "Atas de governança", "Registros de decisões e planos de ação regionais."),
    ]
    add_table(doc, ["Anexo", "Material sugerido", "Finalidade"], rows, [900, 2950, 5510], font_size=8.4)


def audit_document(doc):
    # Minimal deterministic audit markers embedded as custom properties are avoided
    # to keep the deliverable clean. This function intentionally remains a
    # builder-side checkpoint for preset-controlled choices.
    assert doc.sections[0].page_width == Inches(8.5)
    assert doc.sections[0].left_margin == Inches(1)


def main():
    doc = Document()
    configure_section(doc)
    configure_styles(doc)

    core = doc.core_properties
    core.title = "Documentação Final do Projeto - Global Service Governance ZohoDesk"
    core.subject = "Encerramento e handover do projeto Global Service Governance - ZohoDesk"
    core.author = "Guilherme Bernardes Alamino"
    core.keywords = "ZohoDesk, Global Service, Governance, Handover, Playbook, KPI, Supabase"
    core.comments = "Documento gerado para encerramento e handover do projeto."
    core.created = datetime(2026, 6, 18, 12, 0, 0)
    core.modified = datetime(2026, 6, 18, 12, 0, 0)

    build_cover(doc)
    build_document_map(doc)
    build_summary(doc)
    build_context_objectives_scope(doc)
    build_method_timeline(doc)
    build_deliverables(doc)
    build_processes(doc)
    build_dashboards_playbook_automation_ai(doc)
    build_risks_lessons_skills(doc)
    build_handover_backlog_recommendations(doc)
    build_annexes(doc)

    audit_document(doc)
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    main()

