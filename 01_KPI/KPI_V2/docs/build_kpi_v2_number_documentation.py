from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
MD_PATH = OUT_DIR / "documentacao-numeros-kpi-v2.md"
DOCX_PATH = OUT_DIR / "documentacao-numeros-kpi-v2.docx"


TITLE = "Documentacao dos numeros - Playbook Global KPIv2"
SUBTITLE = "Daily, Weekly e Monthly | Fonte, extracao, tratamentos, filtros e formulas"
GENERATED_NOTE = "Gerado em 2026-06-08 a partir do codigo local do Playbook e do projeto Supabase hqaxpbnduupjdhuuwpmg."


LINEAGE = [
    ("Sistema fonte", "Zoho Desk",
     "Tickets, detalhes, metricas, departamentos, agentes, contatos e historico de status sao buscados via API Zoho Desk."),
    ("Ingestao", "Edge Function sync-tickets-v0",
     "Autentica no Zoho, pagina por departamentos, busca detalhes e metricas, converte datas/numeros e grava no Supabase."),
    ("Armazenamento bruto/tratado", "Supabase public.zoho_*",
     "Tabelas principais: zoho_tickets, zoho_agents, zoho_departments, zoho_contacts, zoho_ticket_status_history, sync_runs e sync_errors."),
    ("Camada operacional", "vw_tickets_operational",
     "Junta tickets com agentes, departamentos e contatos; remove is_deleted=true; expõe campos operacionais e flags de qualidade."),
    ("Camada BI canonica", "vw_tickets_bi_base",
     "Normaliza regiao, prioridade, grupo operacional, MTFC, MTTS, aging, metas e status de SLA."),
    ("Views de dashboard", "vw_dashboard_bi_*",
     "Agregam KPIs, backlog e resumos por regiao/status/prioridade a partir da camada BI."),
    ("API de leitura", "Edge Function dashboard-read",
     "Endpoint GET /functions/v1/dashboard-read?type=... com JWT; roteia tipos bi-* para views Supabase."),
    ("Frontend", "01_KPI/KPI_V2/*.html e js/*.js",
     "Aplica cache de 5 minutos, filtros de tela, exclusoes manuais e calculos finais exibidos nas abas."),
]


GLOBAL_RULES = [
    ("Exclusao de deletados",
     "vw_tickets_operational filtra zoho_tickets.is_deleted = false. A funcao de sync possui deletion sweep para marcar tickets removidos no Zoho."),
    ("Exclusao de agentes/parceiros",
     "vw_tickets_bi_base marca is_weekly_report_filter_included=false para nomes contendo Vitor ou Unicorn e para emails Unicorn cadastrados. O frontend tambem remove tickets de dominios de contato unicorndenmart.com e webpeak.com.br."),
    ("Tickets excluidos manualmente",
     "business-rules.js remove ticket_number 220822, 236429 e 236430 antes dos calculos no frontend."),
    ("Regiao canonica",
     "vw_tickets_bi_base tenta region, pais e regiao; se faltarem, usa fallback por email do agente. Mapeia Brasil, Argentina, Mexico, USA, LATAM e ROW; caso contrario, Nao classificado."),
    ("Prioridade canonica",
     "Urgente/P1, Alta/P2, Media/P3, Baixa/P4 e Muito Baixa/P5 sao padronizadas em priority_standard. Prioridade vazia vira Sem prioridade."),
    ("Grupo operacional do agente",
     "Instalacao por categoria com 'instala' ou agente Geovana; Sem dono quando sem agente, Alliage, Norberto ou anomura@prexion.com; Terceiro para Danielly/Contato; Especialista para Camila/Ademar; demais Suporte geral."),
    ("Filtro de inclusao BI",
     "A base usada pelo dashboard considera apenas is_weekly_report_filter_included=true. Na pratica entram Especialista, Sem dono e Suporte geral; ficam fora Instalacao, Terceiro, Unicorn e Vitor."),
    ("MTFC",
     "Vem da metrica Zoho firstResponseTime. A sync converte texto HH:MM para horas e minutos. A BI usa mtfc_minutos/60 quando existe, senao mtfc_horas."),
    ("MTTS",
     "Na BI, mtts_dias_bi = closed_time - created_time em dias. No frontend, quando mtts direto nao existe, ha fallback para resolution_horas/24."),
    ("Aging de backlog",
     "Para tickets abertos, aging_backlog_dias = floor((now - created_time) / 86400). Daily usa esse valor para cards, graficos e tabela."),
    ("Metas MTFC",
     "Urgente 1h, Alta 2h, Media 3h, Baixa 5h, Muito Baixa 6h. Quando prioridade nao encaixa, pode cair no default frontend de 3h em alguns pontos."),
    ("Metas MTTS",
     "Brasil 4 dias, Argentina 5 dias, Mexico 6 dias, LATAM 6 dias, USA 10 dias, ROW 10 dias. No frontend, alguns fallbacks usam 7 dias ou 168 horas."),
    ("Elegibilidade SLA",
     "SQL: ticket fechado, com created_time, regiao valida, meta MTTS e meta MTFC. Frontend computeMetrics exige ticket fechado com MTFC numerico e resolution_horas/mtts disponivel."),
    ("Status SLA SQL",
     "sla_status_bi = Fora SLA quando MTFC excede meta ou closed_time excede created_time + meta MTTS. Caso contrario, Dentro SLA. Sem dados minimos vira Nao elegivel."),
]


ENDPOINTS = [
    ("bi-tickets",
     "vw_tickets_bi_base",
     "Retorna ate 10.000 linhas, filtradas por is_weekly_report_filter_included=true, ordenadas por created_time desc. E a base do Weekly e Monthly."),
    ("bi-backlog",
     "vw_dashboard_bi_backlog",
     "Retorna ate 500 tickets abertos, filtrados por is_weekly_report_filter_included=true, is_open=true e status nao fechado/resolvido. Ordena por aging_backlog_dias desc. E a base do Daily e do backlog canonico no Weekly."),
    ("bi-backlog-kpis",
     "vw_dashboard_bi_backlog_kpis",
     "Agrega backlog_total, sem_primeira_resposta, aging medio, percentuais acima de 7/30 dias e last_data_update. Esta disponivel na API, mas as tres abas atuais calculam os cards no frontend."),
    ("bi-kpis",
     "vw_dashboard_bi_kpis",
     "Agregado geral da camada BI. Usado no hub como teste de conexao; nao alimenta diretamente os cards Daily/Weekly/Monthly atuais."),
    ("bi-region-summary",
     "vw_dashboard_bi_region_summary",
     "Resumo por regiao disponivel para consumo BI; nao e a fonte direta das tres abas atuais."),
    ("bi-summary",
     "vw_dashboard_bi_summary",
     "Resumo por regiao, departamento, prioridade, status e grupo operacional; disponivel para analises auxiliares."),
    ("sync-health",
     "sync_runs, sync_errors, zoho_tickets, vw_tickets_bi_base",
     "Calcula saude da sincronizacao, tickets atualizados em 24h, erros em 24h, frescor e SLA fechado dos ultimos 7 dias."),
]


CANONICAL_FIELDS = [
    ("ticket_id / ticket_number", "Identificador Zoho e numero exibido.", "zoho_tickets.id e ticketNumber."),
    ("created_time", "Data/hora de criacao do ticket.", "Zoho ticket.createdTime; convertida para timestamp."),
    ("closed_time", "Data/hora de fechamento/resolucao.", "Zoho ticket.closedTime; nulo significa aberto na maior parte da UI."),
    ("department_name", "Nome do departamento.", "Join zoho_tickets.department_id = zoho_departments.id."),
    ("agent_name / agent_email", "Nome e email do agente.", "Join zoho_tickets.assignee_id = zoho_agents.id."),
    ("contact_name / contact_email", "Nome e email do solicitante.", "Join zoho_tickets.contact_id = zoho_contacts.id."),
    ("produtos / marca_produto", "Produto/marca usados em filtros e tabelas.", "Campos customizados cf_produtos e cf_marca_do_produto."),
    ("tipo_atendimento", "Tipo de atendimento usado no filtro Type/Tipo.", "Campo customizado cf_tipo_de_atendimento."),
    ("categoria", "Categoria usada para grupo operacional e tabelas.", "Campo customizado cf_categoria."),
    ("regiao_grupo", "Regiao consolidada para dashboards.", "vw_tickets_bi_base a partir de region/pais/regiao e fallback por agent_email."),
    ("priority_standard", "Prioridade consolidada.", "vw_tickets_bi_base padroniza P1/P2/P3/P4/P5 e equivalentes."),
    ("mtfc_horas_bi", "Tempo ate primeira resposta, em horas.", "Metrica Zoho firstResponseTime convertida em sync e normalizada na BI."),
    ("mtts_dias_bi", "Tempo total ate solucao/fechamento, em dias.", "closed_time - created_time em vw_tickets_bi_base."),
    ("aging_backlog_dias", "Idade do ticket aberto em dias inteiros.", "floor(now - created_time) para tickets sem closed_time."),
    ("is_open / is_closed", "Flags de abertura/fechamento.", "closed_time IS NULL / IS NOT NULL."),
    ("meta_mtfc_horas", "Meta de primeira resposta por prioridade.", "Urgente 1h, Alta 2h, Media 3h, Baixa 5h, Muito Baixa 6h."),
    ("meta_mtts_dias", "Meta de solucao por regiao.", "Brasil 4d, Argentina 5d, Mexico/LATAM 6d, USA/ROW 10d."),
    ("sla_status_bi", "Classificacao SQL do SLA.", "Dentro SLA, Fora SLA ou Nao elegivel conforme metas e elegibilidade."),
    ("grupo_operacional_agente", "Grupo operacional usado nos filtros.", "Instalacao, Terceiro, Especialista, Sem dono ou Suporte geral."),
]


DAILY_METRICS = [
    ("Backlog atual", "count(rows)", "Linhas de bi-backlog apos filtros de tela.", "Regiao, prioridade, departamento, status e grupo do agente."),
    ("Sem primeira resposta", "count(rows where mtfc_horas_bi/mtfc_horas nao e numerico)", "Funcao hasFirstResponse(row).", "Mesmos filtros do Daily."),
    ("> 10 dias", "count(rows where aging_backlog_dias > 10)", "agingDays(row).", "Mesmos filtros do Daily."),
    ("Dependencia externa", "count(rows where status normalizado contem aguardando/waiting/terceiro/third/peca/cliente/customer/visita)", "isDependencyStatusText(status).", "Mesmos filtros do Daily."),
    ("P1/P2 abertos", "count(rows where prioridade normalizada e Urgente/Alta/P1/P2/urgent/high)", "isP12(row).", "Mesmos filtros do Daily."),
    ("Aging medio", "avg(aging_backlog_dias)", "Media simples sobre os tickets abertos filtrados.", "Mesmos filtros do Daily."),
    ("Fila critica", "Mesmo numero de > 10 dias", "Decision grid reutiliza over10.", "Mesmos filtros do Daily."),
    ("Primeira resposta", "Mesmo numero de Sem primeira resposta", "Decision grid reutiliza noResponse.", "Mesmos filtros do Daily."),
    ("Dependencias", "Mesmo numero de Dependencia externa", "Decision grid reutiliza dependencies.", "Mesmos filtros do Daily."),
    ("Risk badge", "critical se P1/P2 > 0; attention se >10 dias >0 ou sem resposta >0; senao ok", "Calculo frontend.", "Mesmos filtros do Daily."),
    ("Alertas Daily", "noResponse, Urgente, aging >15 e P1/P2 sem primeira resposta", "renderAlerts(rows).", "Mesmos filtros do Daily."),
    ("Backlog por Status", "count por status, top 10", "groupCounts(rows, status).", "Mesmos filtros do Daily."),
    ("Backlog por Regiao", "count por regiao separado em com/sem primeira resposta", "getTicketRegionGroup + hasFirstResponse.", "Mesmos filtros do Daily."),
    ("Aging Buckets", "count por aging_bucket SQL; fallback floor(agingDays)d", "vw_dashboard_bi_backlog.aging_bucket.", "Mesmos filtros do Daily."),
    ("Backlog por Prioridade", "count por priority_standard/priority, top 10", "groupCounts(rows, prioridade).", "Mesmos filtros do Daily."),
    ("Tabela Tickets em Aberto", "total filtrado, pagina atual e colunas MTFC/Aging/Acao", "Lista rows ordenada por score de risco ou coluna selecionada.", "Busca por ticket_number e ordenacao local."),
]


WEEKLY_METRICS = [
    ("Periodo semanal", "Domingo a sabado; ultimas 8 semanas no seletor", "buildWeekOptions usa o domingo mais recente como ancora.", "A semana selecionada define dateFrom/dateTo."),
    ("SLA Compliance", "withinSla / eligible * 100", "computeMetrics(currentClosedWeek). Eligible exige MTFC e MTTS/resolution numericos.", "Filtros dimensionais + tickets fechados na semana."),
    ("Fechados", "count(tickets with closed_time dentro da semana)", "currentClosedWeek.length.", "Filtros dimensionais + periodo semanal."),
    ("Criados", "count(tickets with created_time dentro da semana)", "currentCreatedWeek.length.", "Filtros dimensionais + periodo semanal."),
    ("MTFC", "avg(mtfc_horas_bi/mtfc_horas) nos fechados", "computeMetrics(currentClosedWeek).avgMtfc.", "Filtros dimensionais + fechados na semana."),
    ("MTTS", "avg(mtts_dias_bi/mtts_dias ou resolution_horas/24) nos fechados", "computeMetrics(currentClosedWeek).avgMtts.", "Filtros dimensionais + fechados na semana."),
    ("Backlog", "count(bi-backlog)", "globalBacklogCount = backlogResult.backlog.length.", "Nao muda com semana nem cross-filter; vem do backlog canonico do Daily."),
    ("Backlog em Aberto", "count(filtered tickets where !closed_time)", "open.length.", "Filtros dimensionais, sem limitar pela semana."),
    ("Delta SLA", "(SLA atual - SLA periodo anterior) / SLA periodo anterior * 100", "deltaPct(curr.slaCompliance, prev.slaCompliance, inverted=false).", "Periodo anterior tem mesma duracao imediatamente antes da semana."),
    ("Delta MTFC", "(MTFC atual - MTFC anterior) / MTFC anterior * 100", "deltaPct(..., inverted=true); reducao e verde.", "Periodo anterior comparavel."),
    ("Delta MTTS", "(MTTS atual - MTTS anterior) / MTTS anterior * 100", "deltaPct(..., inverted=true); reducao e verde.", "Periodo anterior comparavel."),
    ("Resumo por Regiao - Fechados", "count fechados por regiao", "computeByRegion(currentClosedWeek).closed.", "Filtro semanal e dimensoes."),
    ("Resumo por Regiao - MTFC", "avg MTFC por regiao", "computeMetrics(data da regiao).avgMtfc.", "Filtro semanal e dimensoes."),
    ("Resumo por Regiao - MTTS", "avg MTTS por regiao", "computeMetrics(data da regiao).avgMtts.", "Filtro semanal e dimensoes."),
    ("Resumo por Regiao - SLA%", "within/eligible por regiao", "computeMetrics(data da regiao).slaCompliance.", "Filtro semanal e dimensoes."),
    ("Status por Regiao", "On target >=80%, attention >=50%, critical <50%", "Badge da tabela regional.", "Baseado no SLA% da regiao."),
    ("Charts por Regiao", "SLA, MTFC, MTTS e Fechados por regiao", "computeByRegion(currentClosedWeek).", "Cross-filter: cada grafico ignora seu proprio eixo ao recalcular."),
    ("Charts por Prioridade", "SLA, MTFC, MTTS e Fechados por prioridade", "computeMetrics(currentClosedWeek filtrado por prioridade).", "Cross-filter por prioridade/regiao."),
    ("Backlog aberto por Status", "count de tickets abertos por status", "open.filter(!closed_time), agrupado por status.", "Filtros dimensionais, sem periodo."),
    ("Backlog aberto por Aging", "count por buckets 0-3, 4-7, 8-15, 16-30, >30 dias", "_aging = now - created_time.", "Filtros dimensionais, sem periodo."),
    ("Backlog aberto por Regiao", "count de abertos por regiao", "getTicketRegionGroup(open).", "Filtros dimensionais, sem periodo."),
    ("Backlog aberto por Prioridade", "count de abertos por prioridade", "PRIORITY_ORDER.", "Filtros dimensionais, sem periodo."),
    ("Follow-up operacional", "top 30 tickets abertos por aging desc", "open.sort(_aging desc).slice(0,30).", "Filtros dimensionais, sem periodo."),
    ("Tabela Tickets Fechados", "tickets fechados da semana, paginados e ordenaveis", "weeklyClosedList = currentClosedWeek.", "Busca por ticket_number e ordenacao local."),
]


MONTHLY_METRICS = [
    ("Mes/Ano", "Meses derivados de closed_time existentes", "groupByMonth(rawData).", "O seletor filter-month-year controla o periodo efetivo."),
    ("SLA Compliance", "withinSla / eligible * 100", "computeMetrics(monthData).", "Filtros dimensionais + tickets fechados no mes selecionado."),
    ("Tickets fechados", "count(monthData)", "regionTicketsInMonth(filtered, Global, monthKey).", "Filtros dimensionais + closed_time no mes."),
    ("MTFC", "avg(mtfc_horas_bi/mtfc_horas) nos fechados do mes", "computeMetrics(monthData).avgMtfc.", "Filtros dimensionais + mes."),
    ("MTTS", "avg(mtts_dias_bi/mtts_dias ou resolution_horas/24) nos fechados do mes", "computeMetrics(monthData).avgMtts.", "Filtros dimensionais + mes."),
    ("Backlog atual", "count(openBacklog)", "regionOpenBacklog(filtered, Global).", "Filtros dimensionais, mas nao limitado ao mes."),
    ("Badge mensal SLA", "On target >=80%, attention >=50%, critical <50%", "monthly-sla-badge.", "Baseado no SLA Compliance mensal."),
    ("SLA por Regiao", "SLA% por regiao com linha meta 80%", "computeMetrics(monthData filtrado por regiao).", "Somente regioes com fechados/elegiveis."),
    ("MTFC por Regiao", "avg MTFC por regiao", "computeMetrics(...).avgMtfc.", "Cores visuais: <=8h verde, <=16h amarelo, >16h vermelho."),
    ("MTTS por Regiao", "avg MTTS por regiao", "computeMetrics(...).avgMtts.", "Cores por meta regional: <=meta verde, <=2x meta amarelo, >2x meta vermelho."),
    ("Backlog por Regiao", "abertos por regiao, separados em com/sem primeira resposta", "openBacklog + hasFirstResponse.", "Backlog atual, nao mes fechado."),
    ("Tabela Tickets Fechados", "tickets fechados no mes, paginados e ordenaveis", "closedTicketsList = monthData.", "Busca por ticket_number e ordenacao local."),
]


HUB_METRICS = [
    ("Atualizados 24h", "count(zoho_tickets where updated_in_supabase_at >= now - 24h)", "sync-health."),
    ("Erros 24h", "count/list sync_errors where created_at >= now - 24h", "sync-health."),
    ("SLA 7d", "within_sla / (within_sla + outside_sla) * 100", "vw_tickets_bi_base fechados nos ultimos 7 dias; usa sla_status_bi."),
    ("Ultima execucao", "finished_at ou started_at do sync_run mais recente; fallback freshness.updated_in_supabase_at", "sync-health."),
    ("Status geral", "attention se ha erro recente, ticket aberto stale ou ultimo run error; senao ok", "sync-health."),
]


FILTERS = [
    ("Daily - Pais/Regiao", "getTicketRegionGroup(row) == filtro", "Atua sobre bi-backlog ja aberto."),
    ("Daily - Prioridade", "getTicketPriority(row) == filtro", "priority_standard ou priority."),
    ("Daily - Departamento", "department_name/department == filtro", "Na tela o label e Departamento."),
    ("Daily - Status", "status == filtro", "O controle usa data-filter product, mas semanticamente filtra status."),
    ("Daily - Grupo do Agente", "grupo_operacional_agente em agentGroups", "Default: Suporte geral, Especialista, Sem dono."),
    ("Weekly - Pais/Regiao", "getTicketRegionGroup(ticket) == filtro", "Aplicado antes de separar criados, fechados e abertos."),
    ("Weekly - Prioridade", "priority_standard/priority == filtro", "Lista ordenada por Urgente, Alta, Media, Baixa, Muito Baixa."),
    ("Weekly - Tipo", "tipo_atendimento == filtro", "O controle data-filter type usa tipo_atendimento."),
    ("Weekly - Produto", "marca_produto == filtro", "O controle data-filter product usa marca_produto."),
    ("Weekly - Semana", "created_time/closed_time dentro de dateFrom/dateTo", "Define KPIs de criados/fechados e periodo anterior para deltas."),
    ("Weekly - Grupo do Agente", "grupo_operacional_agente em agentGroups", "Default: Suporte geral, Especialista, Sem dono."),
    ("Weekly - Departamento por botao", "Grava FilterState.department", "Observacao: applyFilters nao le department; no codigo atual este botao nao altera os numeros."),
    ("Monthly - Pais/Regiao", "getTicketRegionGroup(ticket) == filtro", "Aplicado antes do recorte mensal."),
    ("Monthly - Prioridade", "priority_standard/priority == filtro", "Aplicado antes do recorte mensal."),
    ("Monthly - Tipo", "tipo_atendimento == filtro", "Aplicado antes do recorte mensal."),
    ("Monthly - Produto", "marca_produto == filtro", "Aplicado antes do recorte mensal."),
    ("Monthly - Mes/Ano", "substring(closed_time, 0, 7) == YYYY-MM", "Este e o periodo efetivo da aba Monthly."),
    ("Monthly - Datas De/Ate", "Campos existem na tela", "Observacao: applyPage zera dateFrom/dateTo; atualmente esses campos nao alteram os numeros."),
    ("Monthly - Grupo do Agente", "grupo_operacional_agente em agentGroups", "No init mensal o default e todos os grupos AGENT_GROUPS."),
]


DEFAULT_FILTER_GUIDE = [
    ("Daily / Backlog",
     "Backlog aberto atual",
     "Default: todas as regiões, todas as prioridades, todos os departamentos, todos os status; grupos selecionados: Suporte geral, Especialista e Sem dono.",
     "Não usa semana, mês ou intervalo de datas. A base já vem somente com tickets abertos e incluídos no filtro BI."),
    ("Weekly",
     "Semana selecionada no topo",
     "Default: semana mais recente do seletor, todas as regiões, todas as prioridades, todos os tipos, todos os produtos; grupos selecionados: Suporte geral, Especialista e Sem dono.",
     "Cards de criados usam created_time na semana. Cards de fechados/SLA/MTFC/MTTS usam closed_time na semana."),
    ("Monthly",
     "Mês/Ano selecionado",
     "Default: mês mais recente disponível no seletor, todas as regiões, todas as prioridades, todos os tipos, todos os produtos; controle de grupo começa com todos os grupos disponíveis na tela.",
     "O período oficial é Mês/Ano. Os campos De/Até aparecem, mas hoje não alteram os números. Na prática, a API bi-tickets já vem limitada aos grupos incluídos no BI."),
    ("Hub / Sync Health",
     "Saúde da sincronização",
     "Default: visão global, sem filtros da tela.",
     "Mostra atualização, erros e SLA dos últimos 7 dias da sincronização, não dos filtros Daily/Weekly/Monthly."),
]


SIMPLE_DAILY_METRICS = [
    ("Backlog atual",
     "Total de tickets que estão abertos agora.",
     "Default: todas as regiões, todas as prioridades, todos os departamentos, todos os status; grupos Suporte geral, Especialista e Sem dono.",
     "Não entra data, semana ou mês. Se o usuário filtrar região, prioridade, departamento, status ou grupo, o total muda."),
    ("Sem primeira resposta",
     "Dentro do backlog atual, quantos tickets ainda não têm MTFC registrado.",
     "Usa exatamente o mesmo recorte do Backlog atual: regiões, prioridades, departamentos, status e grupos selecionados.",
     "Não olha tickets fechados. Não usa mês/semana. É um indicador de pendência de atendimento inicial."),
    ("> 10 dias",
     "Dentro do backlog atual, quantos tickets abertos estão com mais de 10 dias desde a criação.",
     "Usa o mesmo default do Daily e respeita os filtros de região, prioridade, departamento, status e grupo.",
     "Não depende de SLA fechado. É aging de ticket aberto."),
    ("Dependência externa",
     "Conta tickets abertos cujo status indica espera por cliente, terceiro, peça, visita ou status similar.",
     "Usa o mesmo default do Daily e respeita os filtros de região, prioridade, departamento, status e grupo.",
     "O número vem do texto do status. Se o status não estiver padronizado, pode ficar fora dessa contagem."),
    ("P1/P2 abertos",
     "Conta tickets abertos com prioridade alta ou urgente.",
     "Default Daily: todas as regiões/departamentos/status; considera as prioridades P1, P2, Urgente, Alta, Urgent ou High.",
     "Se o filtro de prioridade estiver em outra prioridade, esse número pode cair para zero."),
    ("Aging médio",
     "Média de idade, em dias, dos tickets que continuam abertos.",
     "Usa o mesmo conjunto filtrado do backlog atual.",
     "Tickets fechados não entram. O valor sobe quando há tickets antigos abertos no recorte."),
    ("Fila crítica",
     "Repete o número de tickets abertos com mais de 10 dias.",
     "Mesmo filtro do Daily: região, prioridade, departamento, status e grupo.",
     "É uma leitura operacional do mesmo cálculo de > 10 dias."),
    ("Primeira resposta",
     "Repete o número de tickets abertos sem primeira resposta.",
     "Mesmo filtro do Daily: região, prioridade, departamento, status e grupo.",
     "É a mesma pendência vista em formato de decisão operacional."),
    ("Dependências",
     "Repete o número de dependências externas abertas.",
     "Mesmo filtro do Daily: região, prioridade, departamento, status e grupo.",
     "É a mesma leitura do status de dependência externa."),
    ("Risk badge",
     "Mostra se o recorte está OK, em atenção ou crítico.",
     "Usa os números filtrados do Daily.",
     "Fica crítico se houver P1/P2 aberto; fica atenção se houver ticket >10 dias ou sem primeira resposta."),
    ("Alertas Daily",
     "Lista tickets que merecem ação: sem resposta, urgente, aging maior que 15 dias ou P1/P2 sem resposta.",
     "Usa o mesmo recorte filtrado do Daily.",
     "É uma lista de exceções; não é um total separado da base."),
    ("Backlog por Status",
     "Distribuição do backlog atual por status.",
     "Mesmo default do Daily; se filtrar status, o gráfico mostra somente o status escolhido.",
     "Não usa período. Mostra os 10 status com maior volume."),
    ("Backlog por Região",
     "Distribuição do backlog atual por região, separando com e sem primeira resposta.",
     "Mesmo default do Daily; se filtrar região, o gráfico mostra somente a região selecionada.",
     "Não usa período. A separação depende de existir MTFC."),
    ("Aging Buckets",
     "Agrupa tickets abertos por faixas de idade.",
     "Mesmo default do Daily.",
     "Não usa SLA. Usa aging_bucket vindo da view ou calcula pelos dias em aberto."),
    ("Backlog por Prioridade",
     "Distribuição do backlog atual por prioridade.",
     "Mesmo default do Daily; se filtrar prioridade, mostra somente a prioridade escolhida.",
     "Não usa período. Mostra as 10 prioridades com maior volume."),
    ("Tabela Tickets em Aberto",
     "Lista os tickets abertos do recorte atual.",
     "Mesmo default do Daily, mais busca por número do ticket e ordenação escolhida na tabela.",
     "Busca e ordenação mudam a visualização; os filtros da tela mudam a base."),
]


SIMPLE_WEEKLY_METRICS = [
    ("SLA Compliance",
     "Percentual de tickets fechados na semana que ficaram dentro das metas de primeira resposta e solução.",
     "Default: semana selecionada, todas as regiões, prioridades, tipos e produtos; grupos Suporte geral, Especialista e Sem dono.",
     "Só entram tickets fechados na semana e elegíveis para SLA. Tickets abertos não entram."),
    ("Fechados",
     "Quantidade de tickets fechados na semana selecionada.",
     "Usa closed_time dentro da semana, mais filtros de região, prioridade, tipo, produto e grupo.",
     "Não conta tickets apenas criados na semana se eles ainda não fecharam."),
    ("Criados",
     "Quantidade de tickets criados na semana selecionada.",
     "Usa created_time dentro da semana, mais filtros de região, prioridade, tipo, produto e grupo.",
     "Pode incluir tickets ainda abertos. Não depende de closed_time."),
    ("MTFC",
     "Tempo médio até a primeira resposta dos tickets fechados na semana.",
     "Usa os tickets fechados na semana após filtros de região, prioridade, tipo, produto e grupo.",
     "Tickets sem MTFC numérico não entram na média."),
    ("MTTS",
     "Tempo médio até a solução/fechamento dos tickets fechados na semana.",
     "Usa os tickets fechados na semana após filtros de região, prioridade, tipo, produto e grupo.",
     "Não mede tickets abertos. Quando falta MTTS direto, usa resolution_horas/24."),
    ("Backlog",
     "Total do backlog canônico atual vindo do endpoint de backlog.",
     "Usa as regras globais do bi-backlog, mas não usa os filtros da tela Weekly.",
     "Não muda com semana, região, prioridade, tipo, produto ou cross-filter do Weekly."),
    ("Backlog em Aberto",
     "Quantidade de tickets abertos dentro dos filtros dimensionais da tela Weekly.",
     "Default: todas as regiões, prioridades, tipos e produtos; grupos Suporte geral, Especialista e Sem dono.",
     "Não é limitado à semana. É backlog aberto atual após os filtros dimensionais."),
    ("Delta SLA",
     "Compara o SLA da semana selecionada contra o período anterior de mesmo tamanho.",
     "Usa os mesmos filtros do SLA Compliance.",
     "Se o período anterior não tiver base válida, o delta fica sem comparação."),
    ("Delta MTFC",
     "Compara o MTFC da semana selecionada contra a semana/período anterior.",
     "Usa os mesmos filtros do MTFC.",
     "Redução é boa; aumento é ruim."),
    ("Delta MTTS",
     "Compara o MTTS da semana selecionada contra a semana/período anterior.",
     "Usa os mesmos filtros do MTTS.",
     "Redução é boa; aumento é ruim."),
    ("Resumo por Região",
     "Mostra fechados, MTFC, MTTS e SLA por região na semana.",
     "Usa semana selecionada + filtros de prioridade, tipo, produto e grupo; cada linha é uma região.",
     "Se uma região estiver selecionada no filtro, o resumo fica restrito a ela."),
    ("Charts por Região",
     "Gráficos semanais de SLA, MTFC, MTTS e fechados por região.",
     "Usa tickets fechados na semana e os filtros da tela.",
     "Ao clicar em uma região, os outros gráficos são filtrados por ela."),
    ("Charts por Prioridade",
     "Gráficos semanais de SLA, MTFC, MTTS e fechados por prioridade.",
     "Usa tickets fechados na semana e os filtros da tela.",
     "Ao clicar em uma prioridade, os outros gráficos são filtrados por ela."),
    ("Backlog aberto por Status",
     "Distribuição dos tickets abertos atuais por status.",
     "Usa filtros de região, prioridade, tipo, produto e grupo.",
     "Não usa a semana. É backlog atual, não backlog criado na semana."),
    ("Backlog aberto por Aging",
     "Distribuição dos tickets abertos atuais por faixa de idade.",
     "Usa filtros de região, prioridade, tipo, produto e grupo.",
     "Não usa a semana. Calcula idade desde created_time até hoje."),
    ("Backlog aberto por Região",
     "Distribuição dos tickets abertos atuais por região.",
     "Usa filtros de prioridade, tipo, produto e grupo; pode respeitar região se ela estiver selecionada.",
     "Não usa a semana."),
    ("Backlog aberto por Prioridade",
     "Distribuição dos tickets abertos atuais por prioridade.",
     "Usa filtros de região, tipo, produto e grupo; pode respeitar prioridade se ela estiver selecionada.",
     "Não usa a semana."),
    ("Follow-up operacional",
     "Lista os 30 tickets abertos mais antigos do recorte.",
     "Usa filtros de região, prioridade, tipo, produto e grupo.",
     "Não usa a semana. Serve para ação operacional no backlog atual."),
    ("Tabela Tickets Fechados",
     "Lista os tickets fechados na semana selecionada.",
     "Usa semana + filtros de região, prioridade, tipo, produto e grupo; também aceita busca por número e ordenação.",
     "Busca/ordenação mudam a lista exibida; os filtros mudam a base do número."),
]


SIMPLE_MONTHLY_METRICS = [
    ("SLA Compliance",
     "Percentual de tickets fechados no mês que ficaram dentro das metas.",
     "Default: mês selecionado, todas as regiões, prioridades, tipos e produtos; grupo conforme controle mensal, lembrando que a API já traz a base BI incluída.",
     "Só entram tickets fechados no mês e elegíveis para SLA."),
    ("Tickets fechados",
     "Quantidade de tickets fechados no mês selecionado.",
     "Usa closed_time no mês, mais filtros de região, prioridade, tipo, produto e grupo.",
     "Não conta tickets apenas criados no mês se não fecharam."),
    ("MTFC",
     "Tempo médio até a primeira resposta dos tickets fechados no mês.",
     "Usa tickets fechados no mês após filtros de região, prioridade, tipo, produto e grupo.",
     "Tickets sem MTFC numérico não entram na média."),
    ("MTTS",
     "Tempo médio até solução/fechamento dos tickets fechados no mês.",
     "Usa tickets fechados no mês após filtros de região, prioridade, tipo, produto e grupo.",
     "Não mede tickets abertos. Pode usar resolution_horas/24 como fallback."),
    ("Backlog atual",
     "Quantidade de tickets abertos atualmente dentro dos filtros da tela Monthly.",
     "Usa filtros de região, prioridade, tipo, produto e grupo.",
     "Não usa o mês selecionado. É backlog aberto atual, não backlog do mês."),
    ("Badge mensal SLA",
     "Classifica o mês como on target, atenção ou crítico conforme SLA.",
     "Usa o mesmo SLA Compliance filtrado do mês.",
     "On target >=80%, atenção >=50%, crítico <50%."),
    ("SLA por Região",
     "Mostra o SLA do mês em cada região.",
     "Usa mês selecionado + filtros de prioridade, tipo, produto e grupo; cada barra é uma região.",
     "Se o filtro de região estiver selecionado, o gráfico fica restrito a ela."),
    ("MTFC por Região",
     "Mostra a média de primeira resposta do mês por região.",
     "Usa mês selecionado + filtros da tela.",
     "Cores são visuais: até 8h bom, até 16h atenção, acima de 16h ruim."),
    ("MTTS por Região",
     "Mostra a média de solução do mês por região.",
     "Usa mês selecionado + filtros da tela.",
     "Compara com a meta de cada região."),
    ("Backlog por Região",
     "Mostra tickets abertos atuais por região, separados em com e sem primeira resposta.",
     "Usa filtros de prioridade, tipo, produto e grupo; pode respeitar região selecionada.",
     "Não usa o mês. É backlog atual."),
    ("Tabela Tickets Fechados",
     "Lista tickets fechados no mês selecionado.",
     "Usa mês + filtros de região, prioridade, tipo, produto e grupo; também aceita busca por número e ordenação.",
     "Busca/ordenação mudam a visualização; filtros e mês mudam a base."),
]


SIMPLE_HUB_METRICS = [
    ("Atualizados 24h",
     "Mostra quantos tickets foram atualizados no Supabase nas últimas 24 horas.",
     "Sem filtros de tela.",
     "É indicador de sincronização, não de atendimento."),
    ("Erros 24h",
     "Mostra quantos erros de sincronização ocorreram nas últimas 24 horas.",
     "Sem filtros de tela.",
     "Ajuda a saber se a carga Zoho -> Supabase teve problema recente."),
    ("SLA 7d",
     "Mostra SLA dos tickets fechados nos últimos 7 dias.",
     "Sem filtros de tela; usa status de SLA calculado na view BI.",
     "Não é o mesmo recorte do Daily, Weekly ou Monthly."),
    ("Última execução",
     "Mostra quando a sincronização rodou pela última vez.",
     "Sem filtros de tela.",
     "Serve para validar frescor dos dados."),
    ("Status geral",
     "Resume a saúde da sincronização.",
     "Sem filtros de tela.",
     "Fica em atenção se houver erro recente, ticket aberto stale ou último run com erro."),
]


NON_EFFECTIVE_FILTERS = [
    ("Weekly - Departamento",
     "Existe botão na tela, mas hoje não muda os números.",
     "O estado FilterState.department é preenchido, porém applyFilters não usa esse campo."),
    ("Monthly - Datas De/Até",
     "Existem campos de data na tela, mas hoje não mudam os números.",
     "A aba Monthly usa somente o seletor Mês/Ano como período oficial."),
    ("Weekly - Backlog card",
     "O card Backlog do Weekly não muda com os filtros do Weekly.",
     "Ele vem direto do bi-backlog global, separado dos cálculos semanais."),
]


GOVERNANCE_NOTES = [
    ("SLA SQL vs SLA dos cards",
     "As tabelas exibem sla_status_bi vindo do SQL. Os cards e graficos Weekly/Monthly recalculam SLA no frontend por computeMetrics. Isso pode divergir em tickets sem MTFC/resolution, porque a elegibilidade nao e identica."),
    ("Filtro de departamento no Weekly",
     "A tela tem botoes de departamento, mas applyFilters nao usa FilterState.department. Hoje o numero nao muda ao clicar nesses botoes."),
    ("Datas no Monthly",
     "Os campos De/Ate aparecem, mas applyPage remove dateFrom/dateTo e usa somente Mes/Ano. Para uma aba de documentacao, explicar que Mes/Ano e o filtro oficial."),
    ("Exclusao Oficina no frontend",
     "fetchDashboardBiBacklog tenta filtrar department_id de Oficina, mas o endpoint bi-backlog nao retorna department_id. A exclusao efetiva deve estar garantida na view/API ou ser corrigida no select."),
    ("Backlog Weekly e Monthly",
     "No Weekly, o card Backlog vem de bi-backlog global e nao muda pela semana. No Monthly, Backlog atual e aberto atual filtrado por dimensoes, nao tickets abertos/criados no mes."),
    ("Limites de API",
     "bi-tickets limita a 10.000 linhas e bi-backlog limita a 500. Hoje a base listada possui menos que esses limites para o uso esperado, mas o limite deve ser monitorado."),
    ("Cache",
     "supabase-client.js usa cache em memoria por 5 minutos por query string. Atualizar/refresh limpa cache apenas quando clearCache e chamado."),
]


SOURCES = [
    "Frontend: 01_KPI/KPI_V2/page2-backlog.html, page4-weekly.html, page5-monthly.html.",
    "Regras frontend: 01_KPI/KPI_V2/js/business-rules.js, filters.js, supabase-client.js.",
    "Configuracao: config/supabase.js -> projeto https://hqaxpbnduupjdhuuwpmg.supabase.co.",
    "Edge Function deployada: dashboard-read version 11, verify_jwt=true.",
    "Edge Function local/deployada: sync-tickets-v0 version 50; codigo local em supabase/functions/sync-tickets-v0/index.ts.",
    "Views Supabase consultadas: vw_tickets_operational, vw_tickets_bi_base, vw_dashboard_bi_backlog, vw_dashboard_bi_kpis, vw_dashboard_bi_region_summary, vw_dashboard_bi_summary, vw_dashboard_bi_backlog_kpis.",
    "Tabelas Supabase principais: zoho_tickets, zoho_agents, zoho_departments, zoho_contacts, zoho_ticket_status_history, sync_runs, sync_errors.",
]


def md_table(headers, rows):
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        out.append("| " + " | ".join(str(cell).replace("\n", "<br>") for cell in row) + " |")
    return "\n".join(out)


def md_metric_blocks(rows):
    out = []
    for name, description, filters, note in rows:
        out.append(f"- **{name}:** {description}")
        out.append(f"  - Default/filtros usados: {filters}")
        out.append(f"  - O que não entra ou merece atenção: {note}")
    return "\n".join(out)


def build_markdown():
    parts = [
        f"# {TITLE}",
        SUBTITLE,
        "",
        GENERATED_NOTE,
        "",
        "## 1. Escopo",
        "Este documento explica como os numeros das tres abas oficiais do KPIv2 sao extraidos, tratados, filtrados e calculados: Daily, Weekly e Monthly. Inclui tambem o painel de saude de sincronizacao do hub, porque ele apresenta numeros operacionais relevantes.",
        "",
        "## 2. Leitura simples dos filtros",
        "Esta secao e a leitura de negocio. Ela mostra o recorte padrao de cada aba e quais filtros entram em cada numero, sem depender de entender o codigo.",
        "",
        "### 2.1 Default por aba",
        md_table(["Aba", "Base do numero", "Default da tela", "Observacao simples"], DEFAULT_FILTER_GUIDE),
        "",
        "### 2.2 Daily - filtros por numero",
        md_metric_blocks(SIMPLE_DAILY_METRICS),
        "",
        "### 2.3 Weekly - filtros por numero",
        md_metric_blocks(SIMPLE_WEEKLY_METRICS),
        "",
        "### 2.4 Monthly - filtros por numero",
        md_metric_blocks(SIMPLE_MONTHLY_METRICS),
        "",
        "### 2.5 Hub / Sync Health - filtros por numero",
        md_metric_blocks(SIMPLE_HUB_METRICS),
        "",
        "### 2.6 Filtros que merecem atencao",
        md_table(["Item", "Explicacao simples", "Motivo tecnico"], NON_EFFECTIVE_FILTERS),
        "",
        "## 3. Linhagem dos dados",
        md_table(["Camada", "Local", "Papel"], LINEAGE),
        "",
        "## 4. Endpoints e views usados",
        md_table(["Type dashboard-read", "Fonte Supabase", "Uso"], ENDPOINTS),
        "",
        "## 5. Regras globais e tratamentos",
        md_table(["Regra", "Tratamento"], GLOBAL_RULES),
        "",
        "## 6. Campos canonicos usados nos numeros",
        md_table(["Campo", "Definicao", "Origem"], CANONICAL_FIELDS),
        "",
        "## 7. Filtros efetivos por aba - referencia tecnica",
        md_table(["Aba / Filtro", "Criterio", "Observacao"], FILTERS),
        "",
        "## 8. Numeros da aba Daily - referencia tecnica",
        md_table(["Numero / visual", "Formula", "Fonte", "Filtros"], DAILY_METRICS),
        "",
        "## 9. Numeros da aba Weekly - referencia tecnica",
        md_table(["Numero / visual", "Formula", "Fonte", "Filtros"], WEEKLY_METRICS),
        "",
        "## 10. Numeros da aba Monthly - referencia tecnica",
        md_table(["Numero / visual", "Formula", "Fonte", "Filtros"], MONTHLY_METRICS),
        "",
        "## 11. Numeros do hub / Sync Health - referencia tecnica",
        md_table(["Numero", "Formula", "Fonte"], HUB_METRICS),
        "",
        "## 12. Pontos de governanca e consistencia",
        md_table(["Ponto", "Impacto"], GOVERNANCE_NOTES),
        "",
        "## 13. Referencias tecnicas",
    ]
    parts.extend([f"- {source}" for source in SOURCES])
    parts.extend([
        "",
        "## 14. Resumo rapido para criar uma aba de documentacao",
        "- Fonte oficial: Zoho Desk, sincronizado para Supabase.",
        "- Base canonica do dashboard: vw_tickets_bi_base.",
        "- Endpoint oficial do frontend: dashboard-read.",
        "- Recorte Daily: backlog aberto atual.",
        "- Recorte Weekly: tickets criados/fechados na semana selecionada mais backlog aberto atual.",
        "- Recorte Monthly: tickets fechados no mes selecionado mais backlog aberto atual.",
        "- SLA dos cards: recalculado no frontend por computeMetrics.",
        "- SLA das tabelas: campo sla_status_bi vindo da view SQL.",
    ])
    return "\n".join(parts) + "\n"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C2CC", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_run(run, size=10, bold=False, color="000000"):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        style_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2)
    else:
        run = p.add_run(text)
        style_run(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        style_run(run)


def add_metric_blocks(doc, rows):
    for name, description, filters, note in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        name_run = p.add_run(f"{name}: ")
        style_run(name_run, bold=True, size=9.3, color="1F4D78")
        desc_run = p.add_run(description)
        style_run(desc_run, size=9.2)

        add_bullets(doc, [
            f"Default/filtros usados: {filters}",
            f"O que não entra ou merece atenção: {note}",
        ])


def add_table(doc, headers, rows, widths=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)

    if widths is None:
        widths = [int(9360 / len(headers))] * len(headers)

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        style_run(run, size=font_size, bold=True, color="0B2545")

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(str(value))
            style_run(run, size=font_size)

    doc.add_paragraph()
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(TITLE)
    style_run(run, size=22, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(SUBTITLE)
    style_run(run, size=12, color="1F4D78")

    p = doc.add_paragraph()
    run = p.add_run(GENERATED_NOTE)
    style_run(run, size=9, color="555555")


def build_docx():
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    doc.add_heading("1. Escopo", level=1)
    add_paragraph(doc, "Este documento explica como os numeros das tres abas oficiais do KPIv2 sao extraidos, tratados, filtrados e calculados: Daily, Weekly e Monthly. Inclui tambem o painel de saude de sincronizacao do hub, porque ele apresenta numeros operacionais relevantes.")

    doc.add_heading("2. Leitura simples dos filtros", level=1)
    add_paragraph(doc, "Esta secao e a leitura de negocio. Ela mostra o recorte padrao de cada aba e quais filtros entram em cada numero, sem depender de entender o codigo.")

    doc.add_heading("2.1 Default por aba", level=2)
    add_table(doc, ["Aba", "Base do numero", "Default da tela", "Observacao simples"], DEFAULT_FILTER_GUIDE, widths=[1600, 1900, 4200, 1660], font_size=7.8)

    doc.add_heading("2.2 Daily - filtros por numero", level=2)
    add_metric_blocks(doc, SIMPLE_DAILY_METRICS)

    doc.add_heading("2.3 Weekly - filtros por numero", level=2)
    add_metric_blocks(doc, SIMPLE_WEEKLY_METRICS)

    doc.add_heading("2.4 Monthly - filtros por numero", level=2)
    add_metric_blocks(doc, SIMPLE_MONTHLY_METRICS)

    doc.add_heading("2.5 Hub / Sync Health - filtros por numero", level=2)
    add_metric_blocks(doc, SIMPLE_HUB_METRICS)

    doc.add_heading("2.6 Filtros que merecem atencao", level=2)
    add_table(doc, ["Item", "Explicacao simples", "Motivo tecnico"], NON_EFFECTIVE_FILTERS, widths=[2300, 3300, 3760], font_size=8.3)

    doc.add_heading("3. Linhagem dos dados", level=1)
    add_table(doc, ["Camada", "Local", "Papel"], LINEAGE, widths=[1700, 2300, 5360])

    doc.add_heading("4. Endpoints e views usados", level=1)
    add_table(doc, ["Type", "Fonte Supabase", "Uso"], ENDPOINTS, widths=[1700, 2600, 5060])

    doc.add_heading("5. Regras globais e tratamentos", level=1)
    add_table(doc, ["Regra", "Tratamento"], GLOBAL_RULES, widths=[2300, 7060], font_size=8.7)

    doc.add_heading("6. Campos canonicos usados nos numeros", level=1)
    add_table(doc, ["Campo", "Definicao", "Origem"], CANONICAL_FIELDS, widths=[2100, 3600, 3660], font_size=8.5)

    doc.add_heading("7. Filtros efetivos por aba - referencia tecnica", level=1)
    add_table(doc, ["Aba / Filtro", "Criterio", "Observacao"], FILTERS, widths=[2300, 3200, 3860], font_size=8.5)

    doc.add_heading("8. Numeros da aba Daily - referencia tecnica", level=1)
    add_table(doc, ["Numero / visual", "Formula", "Fonte", "Filtros"], DAILY_METRICS, widths=[2200, 3100, 2400, 1660], font_size=8)

    doc.add_heading("9. Numeros da aba Weekly - referencia tecnica", level=1)
    add_table(doc, ["Numero / visual", "Formula", "Fonte", "Filtros"], WEEKLY_METRICS, widths=[2100, 3100, 2500, 1660], font_size=7.8)

    doc.add_heading("10. Numeros da aba Monthly - referencia tecnica", level=1)
    add_table(doc, ["Numero / visual", "Formula", "Fonte", "Filtros"], MONTHLY_METRICS, widths=[2200, 3100, 2400, 1660], font_size=8)

    doc.add_heading("11. Numeros do hub / Sync Health - referencia tecnica", level=1)
    add_table(doc, ["Numero", "Formula", "Fonte"], HUB_METRICS, widths=[2300, 4700, 2360], font_size=8.5)

    doc.add_heading("12. Pontos de governanca e consistencia", level=1)
    add_table(doc, ["Ponto", "Impacto"], GOVERNANCE_NOTES, widths=[2600, 6760], font_size=8.7)

    doc.add_heading("13. Referencias tecnicas", level=1)
    add_bullets(doc, SOURCES)

    doc.add_heading("14. Resumo rapido para criar uma aba de documentacao", level=1)
    add_bullets(doc, [
        "Fonte oficial: Zoho Desk, sincronizado para Supabase.",
        "Base canonica do dashboard: vw_tickets_bi_base.",
        "Endpoint oficial do frontend: dashboard-read.",
        "Recorte Daily: backlog aberto atual.",
        "Recorte Weekly: tickets criados/fechados na semana selecionada mais backlog aberto atual.",
        "Recorte Monthly: tickets fechados no mes selecionado mais backlog aberto atual.",
        "SLA dos cards: recalculado no frontend por computeMetrics.",
        "SLA das tabelas: campo sla_status_bi vindo da view SQL.",
    ])

    doc.save(DOCX_PATH)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    MD_PATH.write_text(build_markdown(), encoding="utf-8")
    build_docx()
    print(MD_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
